"""Synchronous document parsers for chat attachments (TXT/MD/DOCX/PDF).

P0 仅文档类、请求内同步解析。复用 pypdf（与 rag/parser.py 同源）做文字型 PDF；
DOCX 走 python-docx 转 Markdown。依赖惰性导入，缺失时给清晰错误。
"""
from __future__ import annotations

import io
from dataclasses import dataclass, field
from typing import Any, Optional

from settings import OCR_CONFIG

from .document_ocr_client import DocumentOcrClient, DocumentOcrError

PARSER_VERSION = "document-p0-1"


@dataclass
class ParseResult:
    content_text: str = ""
    structured: dict[str, Any] = field(default_factory=dict)
    char_count: int = 0
    page_count: Optional[int] = None


def _decode_text(data: bytes) -> str:
    """UTF-8 优先、常见中文编码兜底；全部失败则忽略无法解码字符。"""
    for enc in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return data.decode(enc).strip()
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore").strip()


class DocumentProcessor:
    supported_extensions: tuple[str, ...] = ()
    # 每种处理器的提取结果契约版本；升级解析逻辑时 bump，避免旧 extraction 不兼容。
    parser_version: str = PARSER_VERSION

    def parse(self, data: bytes, filename: str) -> ParseResult:
        raise NotImplementedError


class TxtProcessor(DocumentProcessor):
    supported_extensions = ("txt", "md")

    def parse(self, data: bytes, filename: str) -> ParseResult:
        text = _decode_text(data)
        return ParseResult(
            content_text=text,
            structured={"filename": filename},
            char_count=len(text),
        )


class DocxProcessor(DocumentProcessor):
    supported_extensions = ("docx",)

    def parse(self, data: bytes, filename: str) -> ParseResult:
        try:
            from docx import Document  # type: ignore
        except ImportError as exc:  # pragma: no cover - 依赖缺失分支
            raise RuntimeError("DOCX 解析需要 python-docx 依赖") from exc

        document = Document(io.BytesIO(data))
        lines: list[str] = []
        for para in document.paragraphs:
            text = (para.text or "").strip()
            if not text:
                continue
            style = (para.style.name or "").lower() if para.style else ""
            if "heading" in style:
                lines.append(f"## {text}")
            elif "list" in style:
                lines.append(f"- {text}")
            else:
                lines.append(text)
        for table in document.tables:
            lines.append("")
            for row in table.rows:
                cells = [(c.text or "").strip().replace("\n", " ") for c in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
        text = "\n".join(lines).strip()
        return ParseResult(
            content_text=text,
            structured={"filename": filename},
            char_count=len(text),
        )


class PdfProcessor(DocumentProcessor):
    supported_extensions = ("pdf",)
    parser_version = "pdf-query-ocr-v1"

    def __init__(
        self,
        *,
        ocr_client: Any = None,
        ocr_enabled: bool | None = None,
        max_ocr_pages: int | None = None,
    ):
        self._ocr_client = ocr_client
        self.ocr_enabled = (
            bool(OCR_CONFIG.get("enabled")) if ocr_enabled is None else bool(ocr_enabled)
        )
        configured_limit = (
            OCR_CONFIG.get("query_pdf_max_pages", 10)
            if max_ocr_pages is None
            else max_ocr_pages
        )
        self.max_ocr_pages = max(int(configured_limit), 0)

    def _client(self) -> Any:
        if self._ocr_client is None:
            self._ocr_client = DocumentOcrClient()
        return self._ocr_client

    def parse(self, data: bytes, filename: str) -> ParseResult:
        try:
            from pypdf import PdfReader  # type: ignore
        except ImportError as exc:  # pragma: no cover - 依赖缺失分支
            raise RuntimeError("PDF 解析需要 pypdf 依赖") from exc

        reader = PdfReader(io.BytesIO(data))
        page_texts: list[str] = []
        for page in reader.pages:
            try:
                text = (page.extract_text() or "").strip()
            except Exception:
                text = ""
            page_texts.append(text)

        ocr_pages = [index for index, text in enumerate(page_texts, start=1) if not text]
        if ocr_pages:
            if not self.ocr_enabled:
                raise DocumentOcrError("OCR_DISABLED", "PDF 包含扫描页，但 OCR 服务未开启")
            if len(ocr_pages) > self.max_ocr_pages:
                raise DocumentOcrError(
                    "OCR_PAGE_LIMIT_EXCEEDED",
                    f"PDF 包含 {len(ocr_pages)} 个扫描页，超过 {self.max_ocr_pages} 页 OCR 上限",
                )
            self._fill_ocr_pages(data, filename, page_texts, ocr_pages)

        body = "\n\n".join(
            f"{{第 {index} 页}}\n{text}"
            for index, text in enumerate(page_texts, start=1)
        )
        ocr_page_set = set(ocr_pages)
        native_pages = [
            index for index in range(1, len(page_texts) + 1) if index not in ocr_page_set
        ]
        return ParseResult(
            content_text=body,
            structured={
                "filename": filename,
                "pages": list(range(1, len(page_texts) + 1)),
                "native_text_pages": native_pages,
                "ocr_pages": ocr_pages,
            },
            char_count=len(body),
            page_count=len(reader.pages),
        )

    def _fill_ocr_pages(
        self,
        data: bytes,
        filename: str,
        page_texts: list[str],
        ocr_pages: list[int],
    ) -> None:
        """Render and OCR every text-less page; any failure aborts the PDF."""
        try:
            import pymupdf
        except ImportError as exc:  # pragma: no cover - required in production
            raise DocumentOcrError("OCR_FAILED", "PDF OCR 需要 PyMuPDF 依赖") from exc

        try:
            pdf = pymupdf.open(stream=data, filetype="pdf")
        except Exception as exc:
            raise DocumentOcrError("OCR_FAILED", "PDF 扫描页渲染失败") from exc

        try:
            for page_number in ocr_pages:
                try:
                    page = pdf.load_page(page_number - 1)
                    pixmap = page.get_pixmap(matrix=pymupdf.Matrix(2, 2), alpha=False)
                    image_bytes = pixmap.tobytes("png")
                except Exception as exc:
                    raise DocumentOcrError(
                        "OCR_FAILED", f"PDF 第 {page_number} 页渲染失败"
                    ) from exc

                result = self._client().ocr_page(
                    image_bytes,
                    "png",
                    filename=filename,
                    page_number=page_number,
                )
                text = str(getattr(result, "text", "") or "").strip()
                if not text:
                    raise DocumentOcrError(
                        "OCR_EMPTY_RESULT", f"PDF 第 {page_number} 页未识别到文字"
                    )
                page_texts[page_number - 1] = text
        finally:
            pdf.close()
