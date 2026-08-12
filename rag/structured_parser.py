"""Structured parsers for DOCX / CSV / XLSX (audit §11 Phase 2).

Phase 2 opens knowledge-base upload and indexing to DOCX, CSV and XLSX.  Each
parser produces the same §6 block sequence contract as TXT/MD/PDF:

- DOCX preserves paragraph/table *document order* by walking the body XML
  (python-docx's ``paragraphs``/``tables`` list separately and would reorder).
- CSV/XLSX form a sheet/table/cell structure: every table becomes an atomic
  ``table`` block whose ``table_data`` carries ``{sheet, table_id, cells,
  row_span, col_span}``.  The chunker converges that into the chunk-level
  ``table`` citation so answers can locate a sheet/table/row/cell.

The parsers reuse the heading registry (``rag/heading_rules.py``) and the
``HeadingStack`` from ``rag/block_parser.py``; no chunker control-flow change
is needed beyond the row-banding dispatch for overlong tables.
"""
from __future__ import annotations

import csv as _csv
import io
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from .block_parser import HeadingStack
from .encodings import decode_text_bytes
from .heading_rules import match_heading
from .schemas import (
    PARSER_NAMES,
    PARSER_VERSIONS,
    PAGE_STATE_NATIVE_TEXT,
    BLOCK_TYPE_HEADING,
    BLOCK_TYPE_LIST,
    BLOCK_TYPE_PARAGRAPH,
    BLOCK_TYPE_TABLE,
    Block,
    ParsedDocument,
    RawDocument,
    _block_id,
)

# Table separators: ``|`` between cells, newline between rows.  This is the
# deterministic rendering used for BM25/dense matching and for display_text.
_CELL_SEP = " | "
_ROW_SEP = "\n"

# Only ASCII printable whitespace is significant in table cells; fold any
# interior whitespace so the rendered grid stays aligned and greppable.
def _cell_text(value: object) -> str:
    text = "" if value is None else str(value)
    return " ".join(text.split())


def render_cells(cells: List[List[object]]) -> str:
    """Render a 2D cell grid to a pipe-joined text block.

    Shared by the parsers and the chunker's table row-banding so every
    rendering of the same grid is byte-identical (stable chunk hashes).
    """
    if not cells:
        return ""
    return _ROW_SEP.join(
        _CELL_SEP.join(_cell_text(cell) for cell in row)
        for row in cells
        if any(_cell_text(cell) for cell in row)
    )


def _table_data(
    sheet: Optional[str],
    table_id: str,
    cells: List[List[object]],
    row_span: Optional[Dict[str, int]] = None,
    col_span: Optional[Dict[str, int]] = None,
) -> Dict[str, object]:
    """Assemble the block-level table structure (audit §6.1.2 block layer).

    ``row_span``/``col_span`` map the top-left cell coordinate (``"r,c"``) to
    the number of rows/columns the merge covers.  DOCX/CSV leave them empty;
    XLSX fills them from the sheet's merged ranges (audit §14: 合并单元格边界).
    """
    return {
        "sheet": sheet,
        "table_id": table_id,
        "cells": [[_cell_text(cell) for cell in row] for row in cells],
        "row_span": row_span or {},
        "col_span": col_span or {},
    }


def _make_table_block(
    *,
    seq: int,
    prefix: Optional[str],
    heading_path: List[str],
    sheet: Optional[str],
    table_id: str,
    cells: List[List[object]],
    row_span: Optional[Dict[str, int]] = None,
    col_span: Optional[Dict[str, int]] = None,
) -> Block:
    return Block(
        block_id=_block_id(None, seq, prefix=prefix),
        block_type=BLOCK_TYPE_TABLE,
        text=render_cells(cells),
        heading_path=list(heading_path),
        table_data=_table_data(sheet, table_id, cells, row_span=row_span, col_span=col_span),
    )


def _derive_structured_title(blocks: List[Block], filename: str) -> str:
    """Document-level title for structured formats: first heading, else stem."""
    for block in blocks:
        if block.block_type == BLOCK_TYPE_HEADING:
            return block.text
    return Path(filename).stem


def _is_list_paragraph(text: str) -> bool:
    stripped = text.lstrip()
    return bool(stripped) and (
        stripped.startswith(("-", "*", "•")) or stripped.startswith(("1.", "1、"))
    )


class DocxParser:
    """Parse a DOCX preserving paragraph/table original order."""

    supported_file_types = ("docx",)
    parser_name = PARSER_NAMES["docx"]
    parser_version = PARSER_VERSIONS["docx"]

    def parse(self, document: RawDocument) -> List[ParsedDocument]:
        try:
            import docx as _docx
        except ImportError as exc:
            raise RuntimeError("DOCX ingestion requires the optional 'python-docx' package.") from exc

        try:
            parsed_doc = _docx.Document(io.BytesIO(document.content))
        except Exception as exc:
            raise ValueError(f"无法解析 DOCX 文档：{exc}") from exc

        blocks: List[Block] = []
        heading_stack = HeadingStack()
        # Map each ``w:p`` element to its python-docx ``Paragraph`` so the
        # document's native heading style can supplement the regex registry
        # (audit P4 keeps regexes in the registry, not the chunker; using the
        # source document's own structure is the DOCX analogue).
        paragraph_by_element = {
            id(paragraph._p): paragraph for paragraph in parsed_doc.paragraphs
        }
        seq = 0
        for child in parsed_doc.element.body.iterchildren():
            if child.tag.endswith("}p"):
                seq += 1
                text = _paragraph_text(child)
                if not text:
                    continue
                paragraph = paragraph_by_element.get(id(child))
                heading = match_heading(text, "docx") or (
                    _docx_native_heading(paragraph) if paragraph is not None else None
                )
                if heading is not None:
                    level, name = heading
                    heading_stack.push(level, name)
                    blocks.append(
                        Block(
                            block_id=_block_id(None, seq),
                            block_type=BLOCK_TYPE_HEADING,
                            text=name,
                            level=level,
                            heading_path=list(heading_stack.current_path()),
                        )
                    )
                    continue
                block_type = BLOCK_TYPE_LIST if _is_list_paragraph(text) else BLOCK_TYPE_PARAGRAPH
                blocks.append(
                    Block(
                        block_id=_block_id(None, seq),
                        block_type=block_type,
                        text=text,
                        heading_path=list(heading_stack.current_path()),
                    )
                )
            elif child.tag.endswith("}tbl"):
                seq += 1
                cells = _docx_table_cells(child, parsed_doc)
                blocks.append(
                    _make_table_block(
                        seq=seq,
                        prefix=None,
                        heading_path=list(heading_stack.current_path()),
                        sheet=None,
                        table_id=f"t{seq}",
                        cells=cells,
                    )
                )

        title = _derive_structured_title(blocks, document.filename)
        text = _render_blocks_text(blocks)
        return [
            ParsedDocument(
                text=text,
                source_path=document.source_path,
                filename=document.filename,
                file_type=document.file_type,
                page_number=None,
                title=title,
                category=_infer_category(Path(document.source_path)),
                metadata=dict(document.metadata),
                blocks=blocks,
                page_terminal_state=PAGE_STATE_NATIVE_TEXT,
                parser_name=self.parser_name,
                parser_version=self.parser_version,
            )
        ]


_DOCX_STYLE_HEADING_RE = re.compile(r"^(?:heading|标题)\s*(\d)", re.IGNORECASE)
_DOCX_AUTO_NUMBER_RE = re.compile(r"^(?:\d+[.、．]\s*|第[一二三四五六七八九十\d]+[章节部分篇]\s*)+")

def _docx_native_heading(paragraph) -> Optional[Tuple[int, str]]:
    """Map a Word heading style to (level, clean_text), or None.

    The regex registry cannot see style-level structure, so paragraphs styled
    "Heading N" / "标题 N" become heading blocks; auto-numbering prefixes
    ("1. 引言", "一、范围") are stripped the way ``match_heading`` strips them.
    """
    style_name = (getattr(getattr(paragraph, "style", None), "name", None)) or ""
    matched = _DOCX_STYLE_HEADING_RE.match(style_name)
    if not matched:
        return None
    level = int(matched.group(1))
    text = (paragraph.text or "").strip()
    if not text:
        return None
    cleaned = _DOCX_AUTO_NUMBER_RE.sub("", text).strip()
    return level, cleaned or text


def _paragraph_text(element) -> str:
    """Extract paragraph text from a ``w:p`` element, handling runs."""
    parts: List[str] = []
    for node in element.iter():
        if node.tag.endswith("}t"):
            parts.append(node.text or "")
        elif node.tag.endswith("}tab"):
            parts.append("\t")
        elif node.tag.endswith("}br"):
            parts.append("\n")
    text = "".join(parts)
    # A paragraph that is only page/column breaks has no searchable content.
    return " ".join(text.split())


def _docx_table_cells(element, parsed_doc) -> List[List[object]]:
    """Extract the full cell grid from a ``w:tbl`` element via python-docx.

    ``parsed_doc`` maps the ``w:tbl`` element back to its ``Table`` object so
    merged-cell text and bounds are available.
    """
    rows: List[List[object]] = []
    for table in parsed_doc.tables:
        if table._tbl is element:
            for row in table.rows:
                rows.append([_cell_text(cell.text) for cell in row.cells])
            break
    return rows


class CsvParser:
    """Parse CSV into a single table block (sheet=None, table_id="csv")."""

    supported_file_types = ("csv",)
    parser_name = PARSER_NAMES["csv"]
    parser_version = PARSER_VERSIONS["csv"]

    def parse(self, document: RawDocument) -> List[ParsedDocument]:
        text = decode_text_bytes(document.content)
        if not text:
            return []
        try:
            reader = _csv.reader(io.StringIO(text))
            cells = [[_cell_text(cell) for cell in row] for row in reader]
        except Exception as exc:
            raise ValueError(f"无法解析 CSV 文档：{exc}") from exc
        cells = [row for row in cells if any(row)]
        if not cells:
            return []

        blocks = [
            _make_table_block(
                seq=1,
                prefix=None,
                heading_path=[],
                sheet=None,
                table_id="csv",
                cells=cells,
            )
        ]
        title = Path(document.filename).stem
        return [
            ParsedDocument(
                text=_render_blocks_text(blocks),
                source_path=document.source_path,
                filename=document.filename,
                file_type=document.file_type,
                page_number=None,
                title=title,
                category=_infer_category(Path(document.source_path)),
                metadata=dict(document.metadata),
                blocks=blocks,
                page_terminal_state=PAGE_STATE_NATIVE_TEXT,
                parser_name=self.parser_name,
                parser_version=self.parser_version,
            )
        ]


class XlsxParser:
    """Parse each XLSX sheet into a table block with sheet-based identity.

    Each sheet is its own ``ParsedDocument`` (analogous to a PDF page) with
    ``location`` ``s{sheet_index}`` so ``chunk_id`` stays sheet-based instead
    of faking a page (audit §6.1.3: location 段泛化).
    """

    supported_file_types = ("xlsx",)
    parser_name = PARSER_NAMES["xlsx"]
    parser_version = PARSER_VERSIONS["xlsx"]

    def parse(self, document: RawDocument) -> List[ParsedDocument]:
        try:
            import openpyxl
        except ImportError as exc:
            raise RuntimeError("XLSX ingestion requires the optional 'openpyxl' package.") from exc

        try:
            workbook = openpyxl.load_workbook(
                io.BytesIO(document.content),
                data_only=True,
                read_only=False,
            )
        except Exception as exc:
            raise ValueError(f"无法解析 XLSX 文档：{exc}") from exc

        parsed: List[ParsedDocument] = []
        for sheet_index, worksheet in enumerate(workbook.worksheets, start=1):
            cells, row_span, col_span = _sheet_cells(worksheet)
            blocks = [
                _make_table_block(
                    seq=1,
                    prefix=f"s{sheet_index}",
                    heading_path=[],
                    sheet=worksheet.title,
                    table_id=f"t{sheet_index}",
                    cells=cells,
                    row_span=row_span,
                    col_span=col_span,
                )
            ]
            parsed.append(
                ParsedDocument(
                    text=_render_blocks_text(blocks),
                    source_path=document.source_path,
                    filename=document.filename,
                    file_type=document.file_type,
                    page_number=None,
                    title=Path(document.filename).stem,
                    category=_infer_category(Path(document.source_path)),
                    metadata=dict(document.metadata),
                    blocks=blocks,
                    page_terminal_state=PAGE_STATE_NATIVE_TEXT,
                    parser_name=self.parser_name,
                    parser_version=self.parser_version,
                    location=f"s{sheet_index}",
                )
            )
        return parsed


def _sheet_cells(worksheet) -> tuple[List[List[object]], Dict[str, int], Dict[str, int]]:
    """Extract the rectangular cell grid, filling merged-cell values.

    ``openpyxl`` only populates the top-left cell of a merged range; fill the
    value into every covered cell so the rendered table does not lose the
    header/merged label (audit §14: 合并单元格边界).  Also returns the span maps
    keyed by the top-left cell coordinate ``"r,c"``.
    """
    merged_map: Dict[tuple, object] = {}
    row_span: Dict[str, int] = {}
    col_span: Dict[str, int] = {}
    if getattr(worksheet, "merged_cells", None):
        for merged in worksheet.merged_cells.ranges:
            value = merged.start_cell.value
            row_span[f"{merged.min_row},{merged.min_col}"] = merged.max_row - merged.min_row + 1
            col_span[f"{merged.min_row},{merged.min_col}"] = merged.max_col - merged.min_col + 1
            for row in range(merged.min_row, merged.max_row + 1):
                for col in range(merged.min_col, merged.max_col + 1):
                    merged_map[(row, col)] = value

    rows: List[List[object]] = []
    for row in worksheet.iter_rows():
        values = []
        for cell in row:
            if cell.coordinate and merged_map.get((cell.row, cell.column), None) is not None:
                values.append(merged_map[(cell.row, cell.column)])
            else:
                values.append(cell.value)
        if any(_cell_text(value) for value in values):
            rows.append(values)
    return rows, row_span, col_span


def _render_blocks_text(blocks: List[Block]) -> str:
    return "\n\n".join(block.text for block in blocks if block.text).strip()


def _infer_category(path: Path) -> str:
    from .document_loader import infer_category

    return infer_category(path)
