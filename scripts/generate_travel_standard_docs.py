# -*- coding: utf-8 -*-
"""生成 Hommey 企业差旅标准测试文档库（.md / .docx / .pdf）。

单一内容源，三种格式导出，文件名携带分类关键字以命中
rag/document_loader.py 的 DEFAULT_CATEGORY_MAPPING。

用法：
    python scripts/generate_travel_standard_docs.py [--out data/documents]
"""
from __future__ import annotations

import argparse
import os
import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUT = PROJECT_ROOT / "data" / "documents"

# ---------------------------------------------------------------------------
# 内容定义：每个文档 = 标题 + 若干块。
# 块类型：("h1"|"h2"|"h3", text) / ("p", text) / ("list", [items]) / ("table", [rows])
# ---------------------------------------------------------------------------

DOCUMENTS = [
    {
        "stem": "policy_authority_差旅管理制度总则",
        "title": "差旅管理制度总则",
        "blocks": [
            ("h2", "一、目的"),
            ("p", "为规范公司差旅行为、控制差旅成本、保障员工出差安全与合规，制定本制度。所有差旅申请、预订、报销均以本制度及其配套专项制度为准。"),
            ("h2", "二、适用范围"),
            ("list", [
                "适用对象：公司全体正式员工、试用期员工及经批准参与差旅的外部顾问。",
                "适用场景：因公前往常驻工作地以外城市（含港澳台及境外）的差旅活动。",
                "不适用：日常通勤、驻外常驻（另行制度）、公司组织的旅游类团建。",
            ]),
            ("h2", "三、术语定义"),
            ("list", [
                "差旅：因公离开常驻工作地，涉及交通、住宿、餐饮等费用的出行。",
                "国内差旅：中国大陆境内的差旅（不含港澳台）。",
                "国际差旅：前往中国大陆以外地区（含港澳台）的差旅。",
                "一类/二类/三类城市：按《差旅标准总表》划分的城市等级。",
            ]),
            ("h2", "四、管理部门与职责"),
            ("table", [
                ["部门", "职责"],
                ["行政部", "制定与修订差旅标准，管理供应商与预订渠道。"],
                ["财务部", "差旅费用审核、报销、预支与票据合规检查。"],
                ["直属上级", "差旅必要性、行程与预算的初审。"],
                ["员工", "如实申报、合规预订、及时报销并保留票据。"],
            ]),
            ("h2", "五、制度层级与冲突处理"),
            ("p", "制度效力从高到低为：法律与安全要求 > 已批准例外 > 核心规则索引 > 国内/国际/报销专项制度 > 操作指南与 FAQ > 城市信息与环保倡议。"),
            ("p", "同层级条款冲突时，以生效日期较晚的版本为准；仍无法判断时返回“未知/需人工确认”，不得自行推断。"),
            ("h2", "六、生效与版本"),
            ("table", [
                ["版本", "生效日期", "修订说明"],
                ["V1.0", "2026-01-01", "首次发布，建立差旅标准基线。"],
                ["V1.1", "2026-06-01", "上调一类城市住宿限额并细化国际差旅标准。"],
            ]),
        ],
    },
    {
        "stem": "travel_standards_差旅标准总表",
        "title": "差旅标准总表（核心索引）",
        "blocks": [
            ("h2", "一、城市等级划分"),
            ("table", [
                ["城市等级", "代表城市"],
                ["一类城市", "北京、上海、广州、深圳"],
                ["二类城市", "杭州、成都、武汉、南京、西安、重庆、天津、苏州、青岛、长沙"],
                ["三类城市", "除一类、二类以外的其他城市"],
            ]),
            ("h2", "二、普通员工综合标准"),
            ("table", [
                ["城市等级", "住宿限额（元/晚）", "餐费补贴（元/天）"],
                ["一类城市", "500", "100"],
                ["二类城市", "400", "80"],
                ["三类城市", "300", "60"],
            ]),
            ("h2", "三、职级上浮系数"),
            ("table", [
                ["职级", "住宿上浮", "餐补上浮"],
                ["普通员工", "0%", "0%"],
                ["部门经理", "20%", "20%"],
                ["总监", "40%", "40%"],
                ["副总经理及以上", "50%", "50%"],
            ]),
            ("p", "实际限额 = 城市等级标准限额 × (1 + 职级上浮系数)。上浮仅适用于本人，不得转让或合并使用。"),
            ("h2", "四、交通标准速查"),
            ("table", [
                ["交通工具", "普通员工", "部门经理及以上"],
                ["飞机", "经济舱", "商务舱（航程超过 4 小时或总监及以上）"],
                ["高铁/动车", "二等座", "一等座"],
                ["普通火车", "硬卧", "软卧"],
            ]),
        ],
    },
    {
        "stem": "travel_standards_交通出行标准",
        "title": "交通出行标准",
        "blocks": [
            ("h2", "一、航空出行"),
            ("p", "普通员工乘坐经济舱；部门经理及以上且单程航程超过 4 小时，或总监及以上，可乘坐商务舱。机票须通过公司指定差旅平台预订，个人渠道购票原则上不予报销。"),
            ("h2", "二、铁路出行"),
            ("p", "高铁与动车普通员工乘坐二等座，部门经理及以上乘坐一等座；副总经理及以上或单程超过 6 小时可乘坐商务座。夜间长途可乘坐软卧。"),
            ("h2", "三、市内交通"),
            ("list", [
                "出租车、网约车：据实报销，单程超过 200 元需在报销单备注事由。",
                "地铁、公交：据实报销，鼓励使用公共交通。",
                "共享单车/电动车：据实报销，单日累计上限 30 元。",
            ]),
            ("h2", "四、机场与火车站接驳"),
            ("p", "往返机场、火车站优先选择地铁、机场快线；深夜航班或携带大件行李可乘坐出租车/网约车，据实报销。"),
        ],
    },
    {
        "stem": "travel_standards_住宿与餐费标准",
        "title": "住宿与餐费标准",
        "blocks": [
            ("h2", "一、住宿标准"),
            ("table", [
                ["城市等级", "普通员工（元/晚）", "部门经理（元/晚）", "总监及以上（元/晚）"],
                ["一类城市", "500", "600", "700"],
                ["二类城市", "400", "480", "560"],
                ["三类城市", "300", "360", "420"],
            ]),
            ("p", "广州属一类城市，普通员工每晚住宿限额 500 元。住宿费按实际入住间夜核算，含税与服务费，不含早餐。"),
            ("h2", "二、住宿要求"),
            ("list", [
                "优先选择公司协议酒店或指定平台推荐酒店。",
                "报销须同时提供住宿发票与入住水单，缺一不可。",
                "同性两人同行原则上合住一间，标准按较高职级者执行。",
                "超出限额部分由个人承担，不得以其他费用名义报销。",
            ]),
            ("h2", "三、餐费补贴"),
            ("table", [
                ["城市等级", "餐费补贴（元/天）"],
                ["一类城市", "100"],
                ["二类城市", "80"],
                ["三类城市", "60"],
            ]),
            ("p", "餐费补贴按自然日计发；出差当天在途不足半天按半日计发，无需餐费票据。商务宴请须单独申请，不占用餐费补贴。"),
        ],
    },
    {
        "stem": "reimbursement_policy_差旅报销制度",
        "title": "差旅报销制度",
        "blocks": [
            ("h2", "一、报销时限"),
            ("p", "员工应在差旅结束后 10 个工作日内提交报销申请；逾期须书面说明并经部门经理审批，逾期超过 30 天原则上不予报销。"),
            ("h2", "二、票据要求"),
            ("list", [
                "住宿：增值税发票（专票/普票均可）+ 入住水单。",
                "交通：行程单或电子发票，机票另附登机牌或航旅纵横截图。",
                "餐饮：餐费补贴无需票据；商务宴请须提供餐饮发票与事由说明。",
                "所有发票抬头须为“XX 科技有限公司”，税号一致。",
            ]),
            ("h2", "三、差旅预支"),
            ("p", "预计差旅费用超过 2000 元的，可提前 3 个工作日申请预支，额度不超过预估费用的 80%。预支须在差旅结束后 10 个工作日内凭票冲销。"),
            ("h2", "四、报销流程"),
            ("list", [
                "员工在报销系统提交申请并上传票据影像。",
                "直属上级审批差旅必要性。",
                "财务部审核票据合规性。",
                "审批通过后 3-5 个工作日打款至工资卡。",
            ]),
            ("h2", "五、不予报销情形"),
            ("list", [
                "无合规票据或票据信息与行程不符的费用。",
                "私人消费、礼品、娱乐性支出。",
                "未按规定渠道预订或未经审批的超标费用。",
                "重复报销、虚报或篡改票据的费用。",
            ]),
        ],
    },
    {
        "stem": "international_travel_国际差旅标准",
        "title": "国际差旅标准",
        "blocks": [
            ("h2", "一、住宿标准（按国家/地区等级）"),
            ("table", [
                ["国家/地区等级", "住宿限额（美元/晚）", "餐费补贴（美元/天）"],
                ["一类（欧美日韩等）", "200", "60"],
                ["二类（东南亚、中东等）", "120", "40"],
                ["三类（其他）", "80", "30"],
            ]),
            ("h2", "二、交通标准"),
            ("p", "国际航班普通员工乘坐经济舱，总监及以上或单程航程超过 8 小时可乘坐公务舱。境外市内交通据实报销，优先使用公共交通。"),
            ("h2", "三、签证与保险"),
            ("p", "因公签证由行政部统一协助办理，费用据实报销。出国差旅须由公司统一购买境外旅行保险，涵盖医疗、意外与行李延误。"),
            ("h2", "四、外币报销"),
            ("p", "境外费用报销以人民币入账，汇率按行程结束后首个工作日的银行中间价折算；大额外汇支出可申请以外币原币报销。"),
        ],
    },
    {
        "stem": "exception_approval_超标准例外审批",
        "title": "超标准例外审批",
        "blocks": [
            ("h2", "一、审批权限"),
            ("table", [
                ["超标幅度", "审批人"],
                ["10% 以内（含）", "部门经理"],
                ["10% - 30%", "分管副总经理"],
                ["超过 30%", "总经理"],
            ]),
            ("h2", "二、审批流程"),
            ("list", [
                "员工在差旅申请或报销时提交超标说明与依据。",
                "按幅度匹配审批人，逐级审批，不得越级或拆分申请规避审批。",
                "审批通过后按实际额度执行；未获批准的，超标部分由个人承担。",
            ]),
            ("h2", "三、所需材料"),
            ("p", "超标申请须附书面说明，包括超标原因、替代方案评估（如更低价格酒店查询截图）与预估金额；因安全或紧急情况超标的，事后 3 个工作日内补报。"),
            ("h2", "四、违规处理"),
            ("p", "未经审批擅自超标、拆分规避审批或虚假说明的，超标费用不予报销，并按公司员工手册追究相应责任。"),
        ],
    },
    {
        "stem": "faq_差旅常见问题",
        "title": "差旅常见问题（FAQ）",
        "blocks": [
            ("h2", "Q1：广州的出差标准是什么？"),
            ("p", "广州属一类城市，普通员工住宿限额为每晚 500 元，餐费补贴为每天 100 元；交通方面，高铁二等座、飞机经济舱据实报销。"),
            ("h2", "Q2：住宿超标了怎么办？"),
            ("p", "超出限额部分原则上由个人承担；确有正当理由的，按《超标准例外审批》提交申请，由对应职级审批人批准后方可报销。"),
            ("h2", "Q3：报销多久能到账？"),
            ("p", "审批通过后 3-5 个工作日打款至工资卡；差旅结束后 10 个工作日内提交报销，逾期可能影响报销。"),
            ("h2", "Q4：国际出差用什么货币报销？"),
            ("p", "境外费用以人民币入账，汇率按行程结束后首个工作日的银行中间价折算；也可申请外币原币报销。"),
            ("h2", "Q5：差旅可以预支吗？"),
            ("p", "预计差旅费用超过 2000 元可提前 3 个工作日申请预支，额度不超过预估费用的 80%，差旅结束后凭票冲销。"),
            ("h2", "Q6：因故退改签产生的费用能报销吗？"),
            ("p", "因公原因导致的退改签费用据实报销；因个人原因产生的，由个人承担，特殊情况按例外审批流程处理。"),
        ],
    },
    {
        "stem": "city_specific_tips_广州差旅提示",
        "title": "广州差旅提示",
        "blocks": [
            ("h2", "一、城市定位与标准"),
            ("p", "广州属一类城市，普通员工住宿限额每晚 500 元，餐费补贴每天 100 元。核心商圈酒店价格较高，建议提前 3-5 天通过公司指定平台预订。"),
            ("h2", "二、住宿区域建议"),
            ("list", [
                "天河区（珠江新城、体育西）：商务出行首选，靠近主要写字楼。",
                "越秀区（北京路、环市东）：老城区，交通便利、餐饮集中。",
                "海珠区（琶洲）：参加展会、会议建议就近入住。",
            ]),
            ("h2", "三、交通指引"),
            ("list", [
                "白云机场：距市区约 30-45 分钟，地铁 3 号线直达市区。",
                "广州南站：高铁主要枢纽，地铁 2/7/22 号线可达。",
                "市内：地铁线网密集，高峰期建议地铁优先于地面交通。",
            ]),
            ("h2", "四、天气提醒"),
            ("p", "广州春夏多雨、有“回南天”返潮，出差建议携带雨具并留意航班/列车延误；夏季高温多雨，注意防暑。"),
        ],
    },
]


# ---------------------------------------------------------------------------
# Markdown 导出
# ---------------------------------------------------------------------------
def render_md(doc: dict) -> str:
    out = [f"# {doc['title']}", ""]
    for block in doc["blocks"]:
        kind = block[0]
        if kind in ("h1", "h2", "h3"):
            level = {"h1": 1, "h2": 2, "h3": 3}[kind]
            out.append(f"{'#' * level} {block[1]}")
            out.append("")
        elif kind == "p":
            out.append(block[1])
            out.append("")
        elif kind == "list":
            for item in block[1]:
                out.append(f"- {item}")
            out.append("")
        elif kind == "table":
            rows = block[1]
            header = rows[0]
            out.append("| " + " | ".join(header) + " |")
            out.append("| " + " | ".join("---" for _ in header) + " |")
            for row in rows[1:]:
                out.append("| " + " | ".join(row) + " |")
            out.append("")
    return "\n".join(out).strip() + "\n"


# ---------------------------------------------------------------------------
# DOCX 导出
# ---------------------------------------------------------------------------
def render_docx(doc: dict, path: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    d = Document()
    d.add_heading(doc["title"], level=1)
    for block in doc["blocks"]:
        kind = block[0]
        if kind in ("h1", "h2", "h3"):
            d.add_heading(block[1], level={"h1": 1, "h2": 2, "h3": 3}[kind])
        elif kind == "p":
            d.add_paragraph(block[1])
        elif kind == "list":
            for item in block[1]:
                d.add_paragraph(item, style="List Bullet")
        elif kind == "table":
            rows = block[1]
            table = d.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                for c_idx, cell in enumerate(row):
                    table.cell(r_idx, c_idx).text = cell
    d.save(str(path))


# ---------------------------------------------------------------------------
# PDF 导出（PyMuPDF + 内嵌思源黑体，保证 pypdf 可抽取中文文本）
# ---------------------------------------------------------------------------
FONT_PATH = r"C:\Windows\Fonts\simhei.ttf"
PAGE_W, PAGE_H = 595.0, 842.0
MARGIN = 72.0
LINE_H_FACTOR = 1.6


class PdfWriter:
    def __init__(self, font_path: str):
        import fitz

        self.fitz = fitz
        self.font = fitz.Font(fontfile=font_path)
        self.doc = fitz.open()
        self.page = self.doc.new_page(width=PAGE_W, height=PAGE_H)
        self.page.insert_font(fontname="simhei", fontfile=font_path)
        self.y = MARGIN
        self.x = MARGIN
        self.maxw = PAGE_W - 2 * MARGIN

    def _new_page(self) -> None:
        self.page = self.doc.new_page(width=PAGE_W, height=PAGE_H)
        self.page.insert_font(fontname="simhei", fontfile=FONT_PATH)
        self.y = MARGIN

    def _ensure(self, needed: float) -> None:
        if self.y + needed > PAGE_H - MARGIN:
            self._new_page()

    def _wrap(self, text: str, fontsize: float, maxw: float) -> list:
        lines = []
        for raw in text.split("\n"):
            raw = raw.rstrip()
            if not raw:
                lines.append("")
                continue
            cur = ""
            for ch in raw:
                if self.font.text_length(cur + ch, fontsize) <= maxw:
                    cur += ch
                else:
                    if cur:
                        lines.append(cur)
                    cur = ch
            lines.append(cur)
        return lines

    def _emit_lines(self, lines: list, fontsize: float, indent: float = 0.0) -> None:
        lh = fontsize * LINE_H_FACTOR
        for line in lines:
            self._ensure(lh)
            if line:
                self.page.insert_text((self.x + indent, self.y), line, fontname="simhei", fontsize=fontsize)
            self.y += lh

    def title(self, text: str) -> None:
        self._ensure(30)
        self.page.insert_text((self.x, self.y + 14), text, fontname="simhei", fontsize=18)
        self.y += 34

    def heading(self, text: str, level: int) -> None:
        sizes = {1: 15, 2: 12.5, 3: 11}
        before = {1: 16, 2: 12, 3: 10}[level]
        self._ensure(before + sizes[level])
        self.y += before
        self.page.insert_text((self.x, self.y), text, fontname="simhei", fontsize=sizes[level])
        self.y += sizes[level] * LINE_H_FACTOR

    def paragraph(self, text: str) -> None:
        lines = self._wrap(text, 10.5, self.maxw)
        self._emit_lines(lines, 10.5)
        self.y += 5

    def list_items(self, items: list) -> None:
        for item in items:
            lines = self._wrap(item, 10.5, self.maxw - 14)
            self._emit_lines(lines, 10.5, indent=14)
            # 前缀项目符号
            self.y -= 0
        self.y += 5

    def table(self, rows: list) -> None:
        ncols = len(rows[0])
        colw = self.maxw / ncols
        pad = 6.0
        fontsize = 9.5
        lh = fontsize * 1.5

        # 先计算每列包裹后的行数与总高度
        wrapped = []
        col_heights = [0.0] * ncols
        for row in rows:
            cells = []
            row_h = 0.0
            for c, cell in enumerate(row):
                lines = self._wrap(cell, fontsize, colw - 2 * pad)
                cells.append(lines)
                row_h = max(row_h, len(lines) * lh + 2 * pad)
                col_heights[c] = max(col_heights[c], row_h)
            wrapped.append(cells)
            # 记录该行高度
            wrapped[-1].append(row_h)

        table_h = sum(r[-1] for r in wrapped)

        # 若整表放不下，另起一页（表很大时按行分页）
        self._ensure(min(table_h, PAGE_H - 2 * MARGIN))

        top = self.y
        cursor = top
        for r_idx, cells in enumerate(wrapped):
            row_h = cells[-1]
            # 行背景（表头浅灰）
            if r_idx == 0:
                rect = self.fitz.Rect(self.x, cursor, self.x + self.maxw, cursor + row_h)
                self.page.draw_rect(rect, fill=(0.92, 0.92, 0.92), color=(0.6, 0.6, 0.6), width=0.5)
            else:
                rect = self.fitz.Rect(self.x, cursor, self.x + self.maxw, cursor + row_h)
                self.page.draw_rect(rect, color=(0.6, 0.6, 0.6), width=0.5)
            # 单元格文字
            for c, lines in enumerate(cells[:-1]):
                cx = self.x + c * colw + pad
                cy = cursor + pad + fontsize * 0.9
                for line in lines:
                    self.page.insert_text((cx, cy), line, fontname="simhei", fontsize=fontsize)
                    cy += lh
            cursor += row_h
        self.y = cursor + 8


def render_pdf(doc: dict, path: Path) -> None:
    writer = PdfWriter(FONT_PATH)
    writer.title(doc["title"])
    for block in doc["blocks"]:
        kind = block[0]
        if kind in ("h1", "h2", "h3"):
            writer.heading(block[1], {"h1": 1, "h2": 2, "h3": 3}[kind])
        elif kind == "p":
            writer.paragraph(block[1])
        elif kind == "list":
            writer.list_items(block[1])
        elif kind == "table":
            writer.table(block[1])
    writer.doc.save(str(path))
    writer.doc.close()


# ---------------------------------------------------------------------------
def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--out", default=str(DEFAULT_OUT))
    parser.add_argument("--formats", default="md,docx,pdf", help="逗号分隔，如 md,docx,pdf")
    args = parser.parse_args()

    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)
    formats = [f.strip().lower() for f in args.formats.split(",") if f.strip()]

    produced = 0
    for doc in DOCUMENTS:
        for fmt in formats:
            path = out_dir / f"{doc['stem']}.{fmt}"
            if fmt == "md":
                path.write_text(render_md(doc), encoding="utf-8")
            elif fmt == "docx":
                render_docx(doc, path)
            elif fmt == "pdf":
                render_pdf(doc, path)
            else:
                print(f"未知格式: {fmt}", file=sys.stderr)
                continue
            produced += 1
            print(f"生成 {path.relative_to(PROJECT_ROOT)}")

    print(f"\n共生成 {produced} 个文件，输出目录：{out_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
