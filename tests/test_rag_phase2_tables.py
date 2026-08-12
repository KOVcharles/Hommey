"""Phase 2 regression tests (audit §11 Phase 2: DOCX/CSV/XLSX + table structure).

Each test maps to a Phase 2 completion criterion:

- DOCX 保留段落/表格原始顺序  → ``test_docx_preserves_paragraph_table_order``
- Word 原生标题样式进入 heading_path → ``test_docx_native_heading_recognized``
- CSV 形成 sheet/table/cell 结构 → ``test_csv_table_data_and_chunk_citation``
- XLSX 每个 sheet 一页、location=s{n} → ``test_xlsx_sheet_is_page_with_sheet_location``
- 合并单元格填满 + row_span/col_span → ``test_xlsx_merged_cells_filled_and_spans_recorded``
- 超长表格按行分带          → ``test_oversized_table_row_banded``
- chunk 级 table 引用进 metadata → ``test_chunk_metadata_emits_table``
- 兼容层 _coerce_chunk 保留 table → ``test_coerce_chunk_preserves_table``
- 上传侧接受新类型            → ``test_webui_accepts_new_upload_types``
"""
from __future__ import annotations

import io

import pytest

from rag.chunker import BlockChunker, token_count
from rag.config import RAGPipelineConfig
from rag.heading_rules import match_heading
from rag.parser import ParserRegistry, UnsupportedFileTypeError
from rag.retriever import _coerce_chunk
from rag.schemas import (
    BLOCK_TYPE_HEADING,
    BLOCK_TYPE_PARAGRAPH,
    BLOCK_TYPE_TABLE,
    RawDocument,
)
from rag.structured_parser import CsvParser, DocxParser, XlsxParser


# ---- fixtures ---------------------------------------------------------------


def _raw(file_type: str, content: bytes, filename: str) -> RawDocument:
    return RawDocument(
        content=content,
        source_path=f"/tmp/{filename}",
        filename=filename,
        file_type=file_type,
        metadata={"document_version": "v1"},
    )


def _csv_doc(text: str = "city,hotel,standard,note\nShanghai,Hilton,900,含早\nBeijing,Crown Plaza,950,含早\n") -> RawDocument:
    return _raw("csv", text.encode("utf-8"), "费率表.csv")


def _docx_bytes() -> bytes:
    import docx

    document = docx.Document()
    document.add_paragraph("出差住宿标准如下。")
    table = document.add_table(rows=2, cols=3)
    for column, header in enumerate(("城市", "限额", "备注")):
        table.rows[0].cells[column].text = header
    table.rows[1].cells[0].text = "北京"
    table.rows[1].cells[1].text = "950"
    table.rows[1].cells[2].text = "含早"
    document.add_paragraph("国际出差标准另见国际差旅章节。")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _docx_with_heading_bytes() -> bytes:
    import docx

    document = docx.Document()
    document.add_heading("住宿标准", level=1)
    document.add_paragraph("出差住宿标准如下。")
    document.add_heading("餐费标准", level=1)
    document.add_paragraph("每日餐费标准如下。")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _xlsx_bytes(sheet_count: int = 1, rows: int = 3) -> bytes:
    import openpyxl

    workbook = openpyxl.Workbook()
    for index in range(1, sheet_count + 1):
        worksheet = workbook.active if index == 1 else workbook.create_sheet()
        worksheet.title = f"Sheet{index}"
        worksheet.append(["城市", "早餐", "午餐", "晚餐"])
        for row in range(rows):
            worksheet.append([f"城市{row}", 50, 100, 100])
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _xlsx_with_merged_bytes() -> bytes:
    import openpyxl

    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.title = "餐费标准"
    worksheet.append(["城市", "早餐", "午餐", "晚餐"])
    worksheet.append(["北京", 50, 100, 100])
    worksheet.merge_cells("A3:B3")
    worksheet["A3"] = "跨两列"
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


# ---- CSV --------------------------------------------------------------------


def test_csv_table_data_and_chunk_citation():
    parsed = CsvParser().parse(_csv_doc())
    assert len(parsed) == 1
    document = parsed[0]
    assert len(document.blocks) == 1
    table_block = document.blocks[0]
    assert table_block.block_type == BLOCK_TYPE_TABLE
    assert table_block.table_data["table_id"] == "csv"
    assert table_block.table_data["sheet"] is None
    assert len(table_block.table_data["cells"]) == 3  # header + 2 data rows

    chunks = BlockChunker().chunk([document])
    assert len(chunks) == 1
    citation = chunks[0].table
    assert citation == {
        "sheet": None,
        "table_id": "csv",
        "row_start": 0,
        "row_end": 2,
        "col_start": 0,
        "col_end": 3,
    }


# ---- DOCX -------------------------------------------------------------------


def test_docx_preserves_paragraph_table_order():
    document = DocxParser().parse(_raw("docx", _docx_bytes(), "制度.docx"))[0]
    # Body order must be [paragraph, table, paragraph], not paragraphs-then-tables.
    assert [block.block_type for block in document.blocks] == [
        BLOCK_TYPE_PARAGRAPH,
        BLOCK_TYPE_TABLE,
        BLOCK_TYPE_PARAGRAPH,
    ]
    assert document.blocks[1].table_data["table_id"] == "t2"
    assert "北京 | 950 | 含早" in document.blocks[1].text


def test_docx_native_heading_recognized():
    document = DocxParser().parse(_raw("docx", _docx_with_heading_bytes(), "制度.docx"))[0]
    headings = [block for block in document.blocks if block.block_type == BLOCK_TYPE_HEADING]
    assert [block.text for block in headings] == ["住宿标准", "餐费标准"]
    # The paragraph that follows the first heading inherits its heading_path (P2).
    paragraph = document.blocks[1]
    assert paragraph.block_type == BLOCK_TYPE_PARAGRAPH
    assert paragraph.heading_path == ["住宿标准"]

    chunks = BlockChunker().chunk([document])
    assert chunks[0].heading_path == ["住宿标准"]
    assert "c1" in chunks[0].chunk_id.split("::")[2]  # section-based location


# ---- XLSX -------------------------------------------------------------------


def test_xlsx_sheet_is_page_with_sheet_location():
    parsed = XlsxParser().parse(_raw("xlsx", _xlsx_bytes(sheet_count=2), "费率.xlsx"))
    assert len(parsed) == 2
    first, second = parsed
    assert first.location == "s1"
    assert second.location == "s2"
    assert first.blocks[0].table_data["sheet"] == "Sheet1"
    assert second.blocks[0].table_data["sheet"] == "Sheet2"
    assert first.blocks[0].table_data["table_id"] == "t1"

    # block ids are sheet-prefixed so chunk_id's location stays sheet-based.
    assert first.blocks[0].block_id == "s1-b1"

    chunks = BlockChunker().chunk([first, second])
    assert chunks[0].chunk_id.split("::")[2] == "s1"
    assert chunks[1].chunk_id.split("::")[2] == "s2"
    assert chunks[0].table["sheet"] == "Sheet1"


def test_xlsx_merged_cells_filled_and_spans_recorded():
    document = XlsxParser().parse(_raw("xlsx", _xlsx_with_merged_bytes(), "餐费.xlsx"))[0]
    table_data = document.blocks[0].table_data
    # Merged A3:B3 value is filled into BOTH covered cells so the rendered grid
    # does not lose the merged label (audit §14: 合并单元格边界).
    cells = table_data["cells"]
    assert "跨两列" in cells[2][0]
    assert "跨两列" in cells[2][1]
    assert table_data["row_span"].get("3,1") == 1
    assert table_data["col_span"].get("3,1") == 2


# ---- row-banding ------------------------------------------------------------


def test_oversized_table_row_banded():
    import openpyxl

    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.title = "全量费率"
    worksheet.append(["城市", "费率", "说明"])
    for index in range(150):
        worksheet.append([f"城市{index}", 900 + index, f"城市{index}的住宿费率为{index + 900}元每晚。"])
    buffer = io.BytesIO()
    workbook.save(buffer)

    parsed = XlsxParser().parse(_raw("xlsx", buffer.getvalue(), "费率.xlsx"))
    chunks = BlockChunker(max_tokens=400).chunk(parsed)
    assert len(chunks) > 1
    # Every band cites its row window; windows tile the original grid.
    windows = [(chunk.table["row_start"], chunk.table["row_end"]) for chunk in chunks]
    assert windows[0][0] == 0
    assert windows[-1][1] == 150  # header row + 150 data rows
    for start, end in windows:
        assert start <= end
    assert all(token_count(chunk.content) <= 400 for chunk in chunks)


# ---- metadata / compatibility -----------------------------------------------


def test_chunk_metadata_emits_table():
    parsed = CsvParser().parse(_csv_doc())
    chunk = BlockChunker().chunk(parsed)[0]
    metadata = chunk.to_metadata()
    assert metadata["table"]["table_id"] == "csv"
    assert metadata["table"]["row_end"] == 2


def test_coerce_chunk_preserves_table():
    metadata = {
        "source_path": "/tmp/费率.csv",
        "filename": "费率.csv",
        "file_type": "csv",
        "table": {"sheet": None, "table_id": "csv", "row_start": 0, "row_end": 2},
    }
    chunk = _coerce_chunk({"content": "a | b", "metadata": metadata}, index=1)
    assert chunk.table == metadata["table"]


# ---- registry / config / upload ---------------------------------------------


def test_parser_registry_serves_structured_types():
    registry = ParserRegistry()
    for file_type in ("docx", "csv", "xlsx"):
        assert file_type in registry.parsers
    assert registry.parser_version_for("docx") == "docx-block-v1"
    assert registry.parser_version_for("xlsx") == "xlsx-block-v1"
    with pytest.raises(UnsupportedFileTypeError):
        registry.parse(_raw("exe", b"not really", "x.exe"))


def test_config_supported_file_types_resolves_new_formats():
    config = RAGPipelineConfig.from_settings({})
    assert config.supported_file_types == ("txt", "md", "pdf", "docx", "csv", "xlsx")
    narrow = RAGPipelineConfig.from_settings({"supported_file_types": "txt,md"})
    assert narrow.supported_file_types == ("txt", "md")


def test_heading_registry_includes_docx():
    assert match_heading("第三章 差旅标准", "docx") is not None
    assert match_heading("一、住宿标准", "docx") is not None


def test_webui_accepts_new_upload_types():
    from webui_new.knowledge_base_service import SUPPORTED_UPLOAD_TYPES

    assert {"docx", "csv", "xlsx"} <= SUPPORTED_UPLOAD_TYPES


def test_validate_content_rejects_renamed_zip_as_docx():
    from webui_new.knowledge_base_service import KnowledgeBaseManagementService
    from webui_new.core.errors import BusinessError

    with pytest.raises(BusinessError):
        KnowledgeBaseManagementService._validate_content("docx", b"%PDF-not-a-zip")
    # A real ZIP magic byte passes the container check.
    KnowledgeBaseManagementService._validate_content("docx", b"PK\x03\x04rest")
