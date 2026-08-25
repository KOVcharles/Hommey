"""多模态附件 P0：解析器与上传校验的纯单元测试（不连数据库）。"""
import io
import zipfile
from types import SimpleNamespace

import pytest

from multimodal import validation
from multimodal.document_processor import (
    DocxProcessor,
    DocumentProcessor,
    ParseResult,
    PdfProcessor,
    TxtProcessor,
)
from multimodal.document_ocr_client import DocumentOcrError
from multimodal.processors import ProcessorRegistry
from multimodal.service import AttachmentService, get_vision_quota
from multimodal.storage import LocalAttachmentStore
from settings import VISION_CONFIG
from webui_new.core.errors import BusinessError


# ── 解析器 ────────────────────────────────────────────────────────


def _docx_bytes() -> bytes:
    pytest.importorskip("docx")
    from docx import Document

    doc = Document()
    doc.add_heading("报销制度", level=1)
    doc.add_paragraph("住宿标准每天 500 元。")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text, table.rows[0].cells[1].text = "项", "金额"
    table.rows[1].cells[0].text, table.rows[1].cells[1].text = "住宿", "500"
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def test_txt_processor_decodes_utf8_and_gbk():
    proc = TxtProcessor()
    utf8 = proc.parse("差旅报销制度".encode("utf-8"), "policy.txt")
    gbk = proc.parse("差旅报销制度".encode("gbk"), "policy.txt")
    assert utf8.content_text == "差旅报销制度"
    assert gbk.content_text == "差旅报销制度"
    assert utf8.char_count == len("差旅报销制度")


def test_docx_processor_extracts_headings_paragraphs_and_tables():
    data = _docx_bytes()
    result = DocxProcessor().parse(data, "policy.docx")
    assert "报销制度" in result.content_text
    assert "住宿标准每天 500 元" in result.content_text
    # 表格被转成 Markdown 行
    assert "住宿" in result.content_text and "500" in result.content_text
    assert validation.detect_kind("policy.docx", data)[0] == "docx"


class _FakeDocumentOcr:
    def __init__(self, text="# OCR 页面", error: Exception | None = None):
        self.text = text
        self.error = error
        self.calls = []

    def ocr_page(self, data, ext, *, filename, page_number):
        self.calls.append((data, ext, filename, page_number))
        if self.error:
            raise self.error
        return SimpleNamespace(text=self.text, model="fake-ocr", confidence=None)


def _blank_pdf_bytes(page_count: int = 1) -> bytes:
    pytest.importorskip("pypdf")
    from pypdf import PdfWriter

    writer = PdfWriter()
    for _ in range(page_count):
        writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def test_pdf_processor_ocrs_textless_page():
    ocr = _FakeDocumentOcr("# 扫描页\n\n住宿费 500 元")
    result = PdfProcessor(ocr_client=ocr, ocr_enabled=True).parse(
        _blank_pdf_bytes(), "blank.pdf"
    )

    assert result.page_count == 1
    assert result.content_text == "{第 1 页}\n# 扫描页\n\n住宿费 500 元"
    assert result.structured["native_text_pages"] == []
    assert result.structured["ocr_pages"] == [1]
    assert ocr.calls[0][2:] == ("blank.pdf", 1)


def test_pdf_processor_keeps_native_text_without_calling_ocr():
    pymupdf = pytest.importorskip("pymupdf")
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "native PDF text")
    data = document.tobytes()
    document.close()
    ocr = _FakeDocumentOcr(error=AssertionError("native page must not use OCR"))

    result = PdfProcessor(ocr_client=ocr, ocr_enabled=True).parse(data, "native.pdf")

    assert "native PDF text" in result.content_text
    assert result.structured["native_text_pages"] == [1]
    assert result.structured["ocr_pages"] == []
    assert ocr.calls == []


def test_pdf_processor_merges_native_and_ocr_pages_in_page_order():
    pymupdf = pytest.importorskip("pymupdf")
    document = pymupdf.open()
    native = document.new_page()
    native.insert_text((72, 72), "first native page")
    document.new_page()
    data = document.tobytes()
    document.close()

    result = PdfProcessor(
        ocr_client=_FakeDocumentOcr("## second OCR page"),
        ocr_enabled=True,
    ).parse(data, "mixed.pdf")

    assert result.content_text.index("{第 1 页}") < result.content_text.index("{第 2 页}")
    assert "first native page" in result.content_text
    assert "## second OCR page" in result.content_text
    assert result.structured["native_text_pages"] == [1]
    assert result.structured["ocr_pages"] == [2]


def test_pdf_processor_rejects_more_than_configured_ocr_pages_before_calling_model():
    ocr = _FakeDocumentOcr()
    with pytest.raises(DocumentOcrError) as exc:
        PdfProcessor(ocr_client=ocr, ocr_enabled=True, max_ocr_pages=10).parse(
            _blank_pdf_bytes(11), "large-scan.pdf"
        )

    assert exc.value.code == "OCR_PAGE_LIMIT_EXCEEDED"
    assert ocr.calls == []


@pytest.mark.parametrize(
    "ocr",
    [
        _FakeDocumentOcr(text=""),
        _FakeDocumentOcr(error=DocumentOcrError("OCR_TIMEOUT", "timeout")),
    ],
)
def test_pdf_processor_fails_whole_document_when_any_ocr_page_fails(ocr):
    with pytest.raises(DocumentOcrError):
        PdfProcessor(ocr_client=ocr, ocr_enabled=True).parse(
            _blank_pdf_bytes(), "failed-scan.pdf"
        )


def test_pdf_upload_surfaces_typed_ocr_error(tmp_path):
    repository = _UploadRepository()
    processor = PdfProcessor(
        ocr_client=_FakeDocumentOcr(
            error=DocumentOcrError("OCR_TIMEOUT", "timeout")
        ),
        ocr_enabled=True,
    )
    service = AttachmentService(
        store=LocalAttachmentStore(str(tmp_path)),
        repository=repository,
        processors=ProcessorRegistry(processors=[processor]),
    )

    result = service.upload(
        user_id="u1",
        filename="scan.pdf",
        content=_blank_pdf_bytes(),
    )

    assert result.status == "failed"
    assert result.error_code == "OCR_TIMEOUT"
    assert repository.status_updates[-1] == ("u1", "failed", "OCR_TIMEOUT")


def test_processor_registry_dispatches_by_extension():
    reg = ProcessorRegistry()
    assert "docx" in reg.supported_extensions()
    result = reg.parse("txt", "hello".encode("utf-8"), "h.txt")
    assert result.content_text == "hello"
    with pytest.raises(ValueError):
        reg.parse("exe", b"x", "h.exe")


# ── 上传校验 ──────────────────────────────────────────────────────

def test_detect_kind_uses_magic_not_extension():
    ext, kind, mime = validation.detect_kind("policy.pdf", b"%PDF-1.4\n%binary")
    assert ext == "pdf" and kind == "document" and mime == "application/pdf"
    # 内容与扩展名不符
    with pytest.raises(BusinessError):
        validation.detect_kind("policy.pdf", b"not a pdf at all" + b"\x00" * 4)


def test_detect_kind_rejects_double_extension_and_unsupported():
    with pytest.raises(BusinessError):
        validation.detect_kind("invoice.pdf.exe", b"MZ")
    with pytest.raises(BusinessError):
        validation.detect_kind("archive.zip", b"PK\x03\x04" + b"\x00" * 10)


def test_detect_kind_rejects_path_traversal():
    with pytest.raises(BusinessError):
        validation.detect_kind("../evil.pdf", b"%PDF-1.4")


def test_docx_validation_rejects_malformed_and_unsafe_archives(monkeypatch):
    with pytest.raises(BusinessError):
        validation.detect_kind("broken.docx", b"PK\x03\x04not-a-zip")

    unsafe = io.BytesIO()
    with zipfile.ZipFile(unsafe, "w") as archive:
        archive.writestr("[Content_Types].xml", "types")
        archive.writestr("word/document.xml", "document")
        archive.writestr("../outside.xml", "unsafe")
    with pytest.raises(BusinessError):
        validation.detect_kind("unsafe.docx", unsafe.getvalue())

    monkeypatch.setitem(validation.ATTACHMENT_CONFIG, "max_archive_uncompressed_bytes", 10)
    with pytest.raises(BusinessError):
        validation.detect_kind("large.docx", _docx_bytes())


def test_validate_size_and_count():
    validation.validate_size(1024)
    with pytest.raises(BusinessError):
        validation.validate_size(10 ** 12)  # 远超上限
    validation.validate_count(3)
    with pytest.raises(BusinessError):
        validation.validate_count(999)


# ── 私有存储与上传状态机 ──────────────────────────────────────────


def test_local_store_rejects_object_keys_outside_private_root(tmp_path):
    store = LocalAttachmentStore(str(tmp_path))
    key = store.save("u1", "att_1", b"hello")
    assert store.load(key) == b"hello"
    with pytest.raises(ValueError):
        store.load("../outside")
    with pytest.raises(FileExistsError):
        store.save("u1", "att_1", b"overwrite")


class _UploadRepository:
    def __init__(self, fail_create: bool = False):
        self.fail_create = fail_create
        self.attachment = None
        self.created_status = None
        self.extraction = None
        self.status_updates = []

    def create(self, attachment):
        if self.fail_create:
            raise RuntimeError("database unavailable")
        self.attachment = attachment
        self.created_status = attachment.status

    def complete_processing(self, extraction, user_id):
        self.extraction = extraction
        self.attachment.status = "ready"
        self.status_updates.append((user_id, "ready", None))

    def update_status(self, attachment_id, user_id, status, error_code=None):
        self.status_updates.append((user_id, status, error_code))

    def get_many(self, attachment_ids, user_id):
        if self.attachment and self.attachment.id in attachment_ids and self.attachment.user_id == user_id:
            return [self.attachment]
        return []

    def get_extraction(self, attachment_id):
        if self.extraction and self.extraction.attachment_id == attachment_id:
            return self.extraction
        return None


def test_upload_persists_processing_then_extraction_and_ready(tmp_path):
    repository = _UploadRepository()
    store = LocalAttachmentStore(str(tmp_path))
    service = AttachmentService(store=store, repository=repository)

    result = service.upload(
        user_id="u1",
        filename="policy.txt",
        content="hotel limit 500".encode(),
        request_id="request-1",
    )

    assert result.status == "ready"
    assert repository.created_status == "processing"
    assert repository.attachment.status == "ready"
    assert repository.attachment.request_id == "request-1"
    assert repository.attachment.expires_at
    assert repository.extraction.content_text == "hotel limit 500"
    assert repository.status_updates == [("u1", "ready", None)]
    assert store.load(repository.attachment.object_key) == b"hotel limit 500"


def test_docx_upload_and_question_preserve_correct_paragraph_only_in_agent_query(tmp_path):
    repository = _UploadRepository()
    service = AttachmentService(
        store=LocalAttachmentStore(str(tmp_path)),
        repository=repository,
    )
    uploaded = service.upload(
        user_id="u1",
        filename="policy.docx",
        content=_docx_bytes(),
    )

    normalized = service.normalize("请总结住宿限制", [uploaded.id], "u1")

    assert "住宿标准每天 500 元" in normalized.agent_query
    assert "住宿标准每天 500 元" not in normalized.display_message
    assert "policy.docx" in normalized.display_message


def test_upload_removes_private_object_when_metadata_create_fails(tmp_path):
    service = AttachmentService(
        store=LocalAttachmentStore(str(tmp_path)),
        repository=_UploadRepository(fail_create=True),
    )

    with pytest.raises(RuntimeError):
        service.upload(user_id="u1", filename="policy.txt", content=b"hello")

    assert not [path for path in tmp_path.rglob("*") if path.is_file()]


# ── 图片模态：视觉识别路径 ────────────────────────────────────────

_PNG_MAGIC = b"\x89PNG\r\n\x1a\n"


class _FakeImageProcessor(DocumentProcessor):
    """打桩的视觉处理器：不真的调用外部模型。"""

    supported_extensions = ("png", "jpg", "jpeg", "webp")
    parser_version = "image-p0-test"

    def __init__(self, result: ParseResult | None = None, error: Exception | None = None):
        self.result = result or ParseResult(content_text="发票金额 520 元", structured={"amount": "520"})
        self.error = error

    def parse(self, data: bytes, filename: str) -> ParseResult:
        if self.error:
            raise self.error
        return self.result


class _QuotaStub:
    def __init__(self, allowed: bool = True):
        self.allowed = allowed
        self.calls = []

    def consume(self, user_id: str) -> bool:
        self.calls.append(user_id)
        return self.allowed


def _image_service(tmp_path, processor, repository):
    registry = ProcessorRegistry(processors=[processor])
    return AttachmentService(
        store=LocalAttachmentStore(str(tmp_path)),
        repository=repository,
        processors=registry,
    )


def test_image_upload_rejected_when_vision_disabled(tmp_path, monkeypatch):
    monkeypatch.setitem(VISION_CONFIG, "enabled", False)  # 不依赖运行环境变量
    service = _image_service(tmp_path, _FakeImageProcessor(), _UploadRepository())

    with pytest.raises(BusinessError) as exc:
        service.upload(user_id="u1", filename="receipt.png", content=_PNG_MAGIC + b"data")

    assert exc.value.code == "VISION_DISABLED"


def test_image_upload_runs_vision_parser_and_consumes_quota(tmp_path, monkeypatch):
    monkeypatch.setitem(VISION_CONFIG, "enabled", True)
    quota = _QuotaStub(allowed=True)
    monkeypatch.setattr("multimodal.service.get_vision_quota", lambda: quota)
    repository = _UploadRepository()
    processor = _FakeImageProcessor()
    service = _image_service(tmp_path, processor, repository)

    result = service.upload(user_id="u1", filename="receipt.png", content=_PNG_MAGIC + b"data")

    assert result.status == "ready"
    assert result.kind == "image"
    assert quota.calls == ["u1"]
    assert repository.extraction.content_text == "发票金额 520 元"
    assert repository.extraction.parser_version == "image-p0-test"


def test_image_upload_exceeds_daily_quota(tmp_path, monkeypatch):
    monkeypatch.setitem(VISION_CONFIG, "enabled", True)
    monkeypatch.setattr("multimodal.service.get_vision_quota", lambda: _QuotaStub(allowed=False))
    service = _image_service(tmp_path, _FakeImageProcessor(), _UploadRepository())

    with pytest.raises(BusinessError) as exc:
        service.upload(user_id="u1", filename="receipt.png", content=_PNG_MAGIC + b"data")

    assert exc.value.code == "VISION_QUOTA_EXCEEDED"


def test_image_parse_failure_maps_to_vision_error_code(tmp_path, monkeypatch):
    # 真实 ImageProcessor 会把 VisionError 归一化为单参数 RuntimeError(code)，
    # 这里复现该形态（str(exc) == code），验证 service._parse_error_code 映射。
    monkeypatch.setitem(VISION_CONFIG, "enabled", True)
    monkeypatch.setattr("multimodal.service.get_vision_quota", lambda: _QuotaStub(allowed=True))
    processor = _FakeImageProcessor(error=RuntimeError("VISION_FAILED"))
    service = _image_service(tmp_path, processor, _UploadRepository())

    result = service.upload(user_id="u1", filename="receipt.png", content=_PNG_MAGIC + b"data")

    assert result.status == "failed"
    assert result.error_code == "VISION_FAILED"
